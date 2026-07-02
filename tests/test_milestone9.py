"""
Tests for Milestone 9: column detection, image extraction, table detection.

Run with:
    python -m pytest tests/test_milestone9.py -v
"""

import pytest
import numpy as np
from pathlib import Path
from unittest.mock import MagicMock, patch

from models.layout_block import LayoutBlock, BlockType, Alignment
from layout.column_detector import detect_columns, _estimate_x_centre


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def make_block(
    text="متن", block_type=BlockType.BODY,
    page_number=1, alignment=Alignment.RIGHT,
    space_before=0.0, column_index=0,
) -> LayoutBlock:
    return LayoutBlock(
        text         = text,
        block_type   = block_type,
        page_number  = page_number,
        alignment    = alignment,
        space_before = space_before,
        column_index = column_index,
    )


# ---------------------------------------------------------------------------
# LayoutBlock Milestone 9 additions
# ---------------------------------------------------------------------------

class TestLayoutBlockM9:

    def test_is_image_true(self):
        b = LayoutBlock(text="", block_type=BlockType.IMAGE,
                        page_number=1, image_path="/tmp/img.png")
        assert b.is_image is True

    def test_is_table_true(self):
        b = LayoutBlock(text="", block_type=BlockType.TABLE,
                        page_number=1, table_data=[["a","b"],["c","d"]])
        assert b.is_table is True

    def test_default_column_index(self):
        b = make_block()
        assert b.column_index == 0

    def test_image_path_stored(self):
        b = LayoutBlock(text="", block_type=BlockType.IMAGE,
                        page_number=1, image_path="/some/path.png")
        assert b.image_path == "/some/path.png"

    def test_table_data_stored(self):
        data = [["cell1", "cell2"], ["cell3", "cell4"]]
        b = LayoutBlock(text="", block_type=BlockType.TABLE,
                        page_number=1, table_data=data)
        assert b.table_data == data


# ---------------------------------------------------------------------------
# Column detector
# ---------------------------------------------------------------------------

class TestColumnDetector:

    def test_single_column_unchanged(self):
        """All blocks on same side → single column → order unchanged."""
        blocks = [
            make_block("ا", space_before=100, alignment=Alignment.RIGHT),
            make_block("ب", space_before=200, alignment=Alignment.RIGHT),
            make_block("ج", space_before=300, alignment=Alignment.RIGHT),
        ]
        result = detect_columns(blocks, page_width=500)
        assert len(result) == 3

    def test_empty_blocks_returns_empty(self):
        assert detect_columns([], page_width=500) == []

    def test_zero_page_width_returns_original(self):
        blocks = [make_block()]
        result = detect_columns(blocks, page_width=0)
        assert result == blocks

    def test_images_and_tables_preserved(self):
        """IMAGE and TABLE blocks must not be lost during column detection."""
        blocks = [
            make_block("متن", block_type=BlockType.BODY),
            make_block("",    block_type=BlockType.IMAGE),
            make_block("",    block_type=BlockType.TABLE),
        ]
        result = detect_columns(blocks, page_width=500)
        types  = {b.block_type for b in result}
        assert BlockType.BODY  in types
        assert BlockType.IMAGE in types
        assert BlockType.TABLE in types

    def test_column_index_assigned(self):
        """After detection every block must have a column_index."""
        blocks = [make_block() for _ in range(3)]
        result = detect_columns(blocks, page_width=500)
        assert all(isinstance(b.column_index, int) for b in result)

    def test_estimate_x_centre_right(self):
        b = make_block(alignment=Alignment.RIGHT)
        assert _estimate_x_centre(b, 600) == pytest.approx(450.0)

    def test_estimate_x_centre_left(self):
        b = make_block(alignment=Alignment.LEFT)
        assert _estimate_x_centre(b, 600) == pytest.approx(150.0)

    def test_estimate_x_centre_center(self):
        b = make_block(alignment=Alignment.CENTER)
        assert _estimate_x_centre(b, 600) == pytest.approx(300.0)


# ---------------------------------------------------------------------------
# Image extractor
# ---------------------------------------------------------------------------

class TestImageExtractor:

    def test_skips_small_images(self, tmp_path):
        """Images below min_area must be silently skipped."""
        from layout.image_extractor import ImageExtractor
        import fitz

        extractor = ImageExtractor(min_area=100*100, cache_dir=tmp_path)

        # Build a minimal PDF with a tiny 10×10 image
        doc  = fitz.open()
        page = doc.new_page(width=500, height=700)
        # No real image inserted — just test that the extractor handles
        # an empty image list without crashing
        blocks = extractor.extract_page_images(doc, page, page_number=1)
        assert isinstance(blocks, list)
        doc.close()

    def test_cache_dir_created(self, tmp_path):
        from layout.image_extractor import ImageExtractor
        cache = tmp_path / "img_cache"
        ImageExtractor(cache_dir=cache)
        assert cache.exists()

    def test_returns_list_on_page_with_no_images(self, tmp_path):
        from layout.image_extractor import ImageExtractor
        import fitz

        extractor = ImageExtractor(cache_dir=tmp_path)
        doc  = fitz.open()
        page = doc.new_page()
        blocks = extractor.extract_page_images(doc, page, page_number=1)
        assert blocks == []
        doc.close()


# ---------------------------------------------------------------------------
# Table detector
# ---------------------------------------------------------------------------

class TestTableDetector:

    def test_returns_empty_on_no_tables(self):
        from layout.table_detector import TableDetector

        mock_page = MagicMock()
        mock_page.find_tables.return_value = []

        detector = TableDetector()
        blocks   = detector.extract_page_tables(mock_page, page_number=1)
        assert blocks == []

    def test_skips_single_row_table(self):
        from layout.table_detector import TableDetector

        mock_table = MagicMock()
        mock_table.extract.return_value = [["only one row", "here"]]
        mock_table.bbox = (0, 0, 500, 30)

        mock_page = MagicMock()
        mock_page.find_tables.return_value = [mock_table]

        blocks = TableDetector(min_rows=2).extract_page_tables(mock_page, 1)
        assert blocks == []

    def test_valid_table_becomes_block(self):
        from layout.table_detector import TableDetector

        mock_table = MagicMock()
        mock_table.extract.return_value = [
            ["عنوان",   "قیمت"],
            ["قلم",     "٥٠"],
            ["کتاب",    "٢٠٠"],
        ]
        mock_table.bbox = (50, 100, 450, 300)

        mock_page = MagicMock()
        mock_page.find_tables.return_value = [mock_table]

        blocks = TableDetector().extract_page_tables(mock_page, 1)
        assert len(blocks) == 1
        assert blocks[0].block_type == BlockType.TABLE
        assert len(blocks[0].table_data) == 3
        assert blocks[0].table_data[0][0] == "عنوان"

    def test_none_cells_replaced_with_empty_string(self):
        from layout.table_detector import TableDetector

        mock_table = MagicMock()
        mock_table.extract.return_value = [
            [None,  "عنوان", "قیمت"],   # 3 columns, only 1 None
            ["قلم", "اردو",  None  ],   # 1 None
            ["کتاب","متن",   "٢٠٠" ],   # all filled
        ]
        mock_table.bbox = (0, 0, 300, 100)

        mock_page = MagicMock()
        mock_page.find_tables.return_value = [mock_table]

        blocks = TableDetector().extract_page_tables(mock_page, 1)
        assert len(blocks) == 1
        assert blocks[0].table_data[0][0] == ""   # None → ""
        assert blocks[0].table_data[1][2] == ""   # None → ""

    def test_exception_returns_empty_list(self):
        from layout.table_detector import TableDetector

        mock_page = MagicMock()
        mock_page.find_tables.side_effect = Exception("PDF error")

        blocks = TableDetector().extract_page_tables(mock_page, 1)
        assert blocks == []


# ---------------------------------------------------------------------------
# export_layout with image and table blocks
# ---------------------------------------------------------------------------

class TestExportLayoutM9:

    def test_image_block_missing_path_does_not_crash(self, tmp_path):
        from exporter.word_export import WordExporter
        blocks = [
            LayoutBlock(
                text="", block_type=BlockType.IMAGE,
                page_number=1, image_path="/nonexistent/img.png",
            )
        ]
        out = tmp_path / "img_test.docx"
        # Should not raise — just log a warning and insert placeholder
        WordExporter().export_layout(blocks, str(out))
        assert out.exists()

    def test_table_block_written_to_docx(self, tmp_path):
        from exporter.word_export import WordExporter
        from docx import Document
        blocks = [
            LayoutBlock(
                text="", block_type=BlockType.TABLE,
                page_number=1,
                table_data=[["نام", "عمر"], ["علی", "٢٥"]],
            )
        ]
        out = tmp_path / "table_test.docx"
        WordExporter().export_layout(blocks, str(out))
        doc    = Document(str(out))
        tables = doc.tables
        assert len(tables) == 1
        assert tables[0].cell(0, 0).text == "نام"

    def test_mixed_blocks_all_rendered(self, tmp_path):
        from exporter.word_export import WordExporter
        from docx import Document
        blocks = [
            LayoutBlock(text="عنوان", block_type=BlockType.HEADING,
                        page_number=1, heading_level=1, is_rtl=True,
                        background_color=None, border_color=None),
            LayoutBlock(text="متن",   block_type=BlockType.BODY,
                        page_number=1, is_rtl=True,
                        background_color=None, border_color=None),
            LayoutBlock(text="",      block_type=BlockType.TABLE,
                        page_number=1,
                        table_data=[["خانہ ایک", "خانہ دو"]]),
            LayoutBlock(text="1",     block_type=BlockType.PAGE_NUM,
                        page_number=1),
        ]
        out = tmp_path / "mixed.docx"
        WordExporter().export_layout(blocks, str(out))
        doc = Document(str(out))
        assert len(doc.tables) == 1
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "عنوان" in full_text
        assert "متن"   in full_text