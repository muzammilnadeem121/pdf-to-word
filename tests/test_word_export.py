"""
Tests for WordExporter.

Run with:
    python -m pytest tests/test_word_export.py -v
"""

import pytest
from pathlib import Path
from docx import Document

from extractor.extractor import ExtractionResult, PageResult
from extractor.scanner import ScanDetector
from models.page_classification import PageClassification, PageType


# ---------------------------------------------------------------------------
# Helper — build a PageClassification without needing a real fitz.Page
# ---------------------------------------------------------------------------

def make_classification(
    page_number: int,
    page_type: PageType,
    char_count: int = 0,
    image_coverage: float = 0.0,
    urdu_char_ratio: float = 0.0,
) -> PageClassification:
    return PageClassification(
        page_number     = page_number,
        page_type       = page_type,
        confidence      = 0.9,
        char_count      = char_count,
        image_coverage  = image_coverage,
        urdu_char_ratio = urdu_char_ratio,
        reason          = "test fixture",
    )


def make_digital_page(page_number: int, text: str) -> PageResult:
    classification = make_classification(
        page_number  = page_number,
        page_type    = PageType.DIGITAL,
        char_count   = len(text.replace(" ", "")),
        urdu_char_ratio = 0.8,
    )
    return PageResult(
        page_number    = page_number,
        classification = classification,
        raw_text       = text,
        char_count     = len(text.replace(" ", "")),
    )


def make_scanned_page(page_number: int) -> PageResult:
    classification = make_classification(
        page_number    = page_number,
        page_type      = PageType.SCANNED,
        image_coverage = 0.95,
    )
    return PageResult(
        page_number    = page_number,
        classification = classification,
        raw_text       = None,
        char_count     = 0,
    )


def make_result(pages: list[PageResult]) -> ExtractionResult:
    digital = sum(1 for p in pages if p.is_digital)
    scanned = sum(1 for p in pages if p.is_scanned)
    mixed   = sum(1 for p in pages if p.is_mixed)
    return ExtractionResult(
        file_path     = "dummy.pdf",
        total_pages   = len(pages),
        digital_pages = digital,
        scanned_pages = scanned,
        mixed_pages   = mixed,
        pages         = pages,
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestWordExporter:

    def test_creates_file(self, tmp_path):
        from exporter.word_export import WordExporter
        result = make_result([make_digital_page(1, "یہ ایک آزمائش ہے")])
        out = tmp_path / "out.docx"
        returned = WordExporter().export(result, str(out))
        assert out.exists()
        assert returned == str(out)

    def test_scanned_placeholder_written(self, tmp_path):
        from exporter.word_export import WordExporter
        result = make_result([make_scanned_page(1)])
        out = tmp_path / "scanned.docx"
        WordExporter().export(result, str(out))
        doc = Document(str(out))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "OCR" in full_text or "Scanned" in full_text.lower() or "اسکین" in full_text

    def test_multipage_has_content(self, tmp_path):
        from exporter.word_export import WordExporter
        result = make_result([
            make_digital_page(1, "پہلا صفحہ"),
            make_digital_page(2, "دوسرا صفحہ"),
        ])
        out = tmp_path / "multi.docx"
        WordExporter().export(result, str(out))
        doc = Document(str(out))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "پہلا" in full_text
        assert "دوسرا" in full_text

    def test_raises_on_empty_result(self, tmp_path):
        from exporter.word_export import WordExporter
        result = make_result([])
        with pytest.raises(ValueError):
            WordExporter().export(result, str(tmp_path / "empty.docx"))

    def test_output_dir_created_automatically(self, tmp_path):
        from exporter.word_export import WordExporter
        result = make_result([make_digital_page(1, "متن")])
        deep_path = tmp_path / "a" / "b" / "c" / "out.docx"
        WordExporter().export(result, str(deep_path))
        assert deep_path.exists()

    def test_mixed_page_writes_text_and_note(self, tmp_path):
        from exporter.word_export import WordExporter
        classification = make_classification(
            page_number    = 1,
            page_type      = PageType.MIXED,
            char_count     = 30,
            image_coverage = 0.5,
        )
        page = PageResult(
            page_number    = 1,
            classification = classification,
            raw_text       = "کچھ متن موجود ہے",
            char_count     = 14,
        )
        result = make_result([page])
        out = tmp_path / "mixed.docx"
        WordExporter().export(result, str(out))
        doc = Document(str(out))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "کچھ متن" in full_text
        assert "OCR" in full_text or "image" in full_text.lower()