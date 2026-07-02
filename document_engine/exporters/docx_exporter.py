"""
Stage 8 — DOCX Exporter
--------------------------
Consumes Document (Stage 7 output) EXCLUSIVELY. Never touches the PDF.
"""

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from document_engine.dom.document import Document as DOMDocument

logger = logging.getLogger(__name__)

_HEADING_SIZES = {1: 20, 2: 16, 3: 13}


class DocxExporter:
    """Exports a Document (DOM) to a .docx file. Consumes DOM only."""

    def __init__(self, font_name: str = "Jameel Noori Nastaleeq") -> None:
        self.font_name = font_name

    def export(self, document: DOMDocument, output_path: str) -> str:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        docx = DocxDocument()
        self._set_document_rtl(docx)

        for page in document.pages:
            for element in sorted(page.elements, key=lambda e: e.sequence):
                self._write_element(docx, element)

        docx.save(str(out))
        logger.info("DOCX exported: %s", out)
        return str(out)

    def _write_element(self, docx: DocxDocument, element) -> None:
        et = element.element_type

        if et == "heading":
            self._write_heading(docx, element)
        elif et == "paragraph":
            self._write_paragraph(docx, element.text)
        elif et == "list":
            for item in element.list_items:
                self._write_paragraph(docx, item, style="List Bullet")
        elif et == "quotation":
            self._write_paragraph(docx, element.text, italic=True, indent=True)
        elif et in ("reference", "footnote"):
            self._write_paragraph(docx, element.text, size=10)
        elif et == "image_caption":
            self._write_paragraph(docx, element.caption_text or "", italic=True, size=10)
        elif et == "table":
            self._write_table(docx, element.table_data)
        else:
            if element.text:
                self._write_paragraph(docx, element.text)

    def _write_heading(self, docx, element) -> None:
        level = max(1, min(element.heading_level or 2, 3))
        para = docx.add_heading(level=level)
        self._set_paragraph_rtl(para)
        run = para.add_run(element.text)
        run.font.name = self.font_name
        run.font.size = Pt(_HEADING_SIZES.get(level, 13))
        run.font.bold = True
        self._set_run_rtl(run)

    def _write_paragraph(
        self, docx, text: str, italic=False, size=12, indent=False, style=None,
    ) -> None:
        if not text.strip():
            return
        para = docx.add_paragraph(style=style) if style else docx.add_paragraph()
        self._set_paragraph_rtl(para)
        if indent:
            para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(size)
        run.font.italic = italic
        self._set_run_rtl(run)

    def _write_table(self, docx, table_data: list[list[str]]) -> None:
        if not table_data:
            return
        rows = len(table_data)
        cols = max(len(r) for r in table_data)
        table = docx.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for r, row in enumerate(table_data):
            for c, cell_text in enumerate(row):
                if c >= cols:
                    break
                cell = table.cell(r, c)
                cell.text = cell_text
                for para in cell.paragraphs:
                    self._set_paragraph_rtl(para)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _set_document_rtl(self, docx) -> None:
        try:
            normal = docx.styles["Normal"]
            pPr = normal.element.get_or_add_pPr()
            pPr.append(OxmlElement("w:bidi"))
        except Exception:
            pass

    def _set_paragraph_rtl(self, para) -> None:
        pPr = para._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:bidi"))
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _set_run_rtl(self, run) -> None:
        rPr = run._r.get_or_add_rPr()
        rPr.append(OxmlElement("w:rtl"))