"""
WordExporter
------------
Converts a structured ExtractionResult into a Microsoft Word (.docx) file.

Responsibilities:
  - Set document-level RTL direction for Urdu text.
  - Write each page's text as properly formatted paragraphs.
  - Insert page breaks between PDF pages.
  - Insert placeholders for scanned pages (OCR fills these in Milestone 6).
  - Apply correct Urdu font and font size.

This module does NOT perform extraction, OCR, or Unicode repair.
It only takes already-extracted text and writes it to a DOCX.
"""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.layout_block import LayoutBlock, BlockType, Alignment
from extractor.extractor import ExtractionResult, PageResult
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement


logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants — override via constructor if needed
# ---------------------------------------------------------------------------

DEFAULT_FONT = "Jameel Noori Nastaleeq"
FALLBACK_FONT = "Noto Nastaliq Urdu"
DEFAULT_FONT_SIZE = Pt(14)   # Nastaliq needs a bit more size to be readable
PLACEHOLDER_COLOR = RGBColor(0x99, 0x99, 0x99)  # grey for scanned placeholders


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

def _make_bidi_element() -> OxmlElement:
    """Create a <w:bidi/> element that sets paragraph direction to RTL."""
    return OxmlElement("w:bidi")


def _make_rtl_element() -> OxmlElement:
    """Create a <w:rtl/> element that sets run direction to RTL."""
    return OxmlElement("w:rtl")


def _set_paragraph_rtl(paragraph) -> None:
    """
    Apply RTL direction to a paragraph.

    We must inject raw XML because python-docx does not expose
    paragraph bidi direction as a first-class property.
    """
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing bidi element first (avoid duplicates on re-runs)
    for existing in pPr.findall(qn("w:bidi")):
        pPr.remove(existing)
    pPr.append(_make_bidi_element())
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_run_rtl(run) -> None:
    """Apply RTL direction to a run (inline text span)."""
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn("w:rtl")):
        rPr.remove(existing)
    rPr.append(_make_rtl_element())


def _set_document_rtl(doc: Document) -> None:
    """
    Set the document default paragraph style to RTL.

    This ensures that any paragraph that doesn't explicitly set a direction
    still renders correctly in Word's RTL mode.
    """
    # Access the document body's <w:body> → <w:sectPr> → <w:bidi>
    # We set it at the styles level instead, which is cleaner.
    try:
        normal_style = doc.styles["Normal"]
        pPr = normal_style.element.get_or_add_pPr()
        for existing in pPr.findall(qn("w:bidi")):
            pPr.remove(existing)
        pPr.append(_make_bidi_element())
    except Exception as exc:
        logger.warning("Could not set document-level RTL: %s", exc)


# ---------------------------------------------------------------------------
# Main exporter class
# ---------------------------------------------------------------------------

class WordExporter:
    """
    Converts an ExtractionResult to a .docx file.

    Parameters
    ----------
    font_name : str
        Urdu font to use. Must be installed on the system where Word opens
        the file. Defaults to Jameel Noori Nastaleeq.
    font_size : Pt
        Font size for body text. Default is 14pt (Nastaliq reads better large).
    """

    def __init__(
        self,
        font_name: str = DEFAULT_FONT,
        font_size: Pt = DEFAULT_FONT_SIZE,
    ) -> None:
        self.font_name = font_name
        self.font_size = font_size

    def export(self, extraction: ExtractionResult, output_path: str) -> str:
        """
        Write a .docx file from an ExtractionResult.

        Parameters
        ----------
        extraction : ExtractionResult
            The structured result from PDFExtractor.
        output_path : str
            Full path where the .docx file should be saved.

        Returns
        -------
        str
            The resolved output path (same as input, useful for chaining).

        Raises
        ------
        ValueError
            If the extraction result has no pages.
        IOError
            If the file cannot be written.
        """
        if not extraction.pages:
            raise ValueError("ExtractionResult has no pages — nothing to export.")

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        _set_document_rtl(doc)

        logger.info(
            "Exporting %d pages to %s", extraction.total_pages, output_path
        )

        for index, page in enumerate(extraction.pages):
            self._write_page(doc, page)

            # Insert a Word page break after every page except the last
            if index < len(extraction.pages) - 1:
                self._add_page_break(doc)

        try:
            doc.save(str(out))
            logger.info("DOCX saved: %s", out)
        except Exception as exc:
            raise IOError(f"Could not save DOCX to {output_path}: {exc}") from exc

        return str(out)

    # ── Private helpers ───────────────────────────────────────────────────────

    def _write_page(self, doc: Document, page: PageResult) -> None:
        if page.error:
            self._write_error_placeholder(doc, page.page_number, page.error)
        elif page.is_scanned:
            if page.ocr_result and not page.ocr_result.is_empty:
                # OCR succeeded — write the recognized text
                self._write_text(doc, page)
            else:
                # OCR not run or returned nothing — show placeholder
                self._write_scanned_placeholder(doc, page.page_number)
        elif page.is_mixed:
            self._write_text(doc, page)
            if not page.ocr_result or page.ocr_result.is_empty:
                self._write_mixed_note(doc, page.page_number)
        else:
            self._write_text(doc, page)

    def _write_mixed_note(self, doc: Document, page_number: int) -> None:
        """Small grey note on mixed pages indicating partial OCR is pending."""
        para = doc.add_paragraph()
        run  = para.add_run(
            f"[ Page {page_number} — Contains images. OCR will be added in a future step. ]"
        )
        run.font.name   = "Arial"
        run.font.size   = Pt(9)
        run.font.color.rgb = PLACEHOLDER_COLOR
        run.font.italic = True

    def _write_text(self, doc: Document, page: PageResult) -> None:
        """
        Write the best available text for a page.
        Uses final_text which merges direct extraction + OCR results.
        """
        text = page.final_text   # ← was page.raw_text
        lines = [line for line in text.splitlines() if line.strip()]

        if not lines:
            return

        for line in lines:
            para = doc.add_paragraph()
            _set_paragraph_rtl(para)
            run = para.add_run(line.strip())
            run.font.name = self.font_name
            run.font.size = self.font_size
            _set_run_rtl(run)

    def _write_scanned_placeholder(self, doc: Document, page_number: int) -> None:
        """
        Insert a visible grey placeholder for a scanned page.

        In Milestone 6, the OCR engine will replace this with real text.
        """
        para = doc.add_paragraph()
        run = para.add_run(
            f"[ صفحہ {page_number} — اسکین شدہ، OCR درکار ہے ]"
            f"\n[ Page {page_number} — Scanned. OCR pending. ]"
        )
        run.font.name = self.font_name
        run.font.size = Pt(11)
        run.font.color.rgb = PLACEHOLDER_COLOR
        run.font.italic = True

    def _write_error_placeholder(
        self, doc: Document, page_number: int, error: str
    ) -> None:
        """Insert a visible placeholder when a page failed to extract."""
        para = doc.add_paragraph()
        run = para.add_run(
            f"[ Page {page_number} — Extraction error: {error} ]"
        )
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # red
        run.font.italic = True

    def _add_page_break(self, doc: Document) -> None:
        """Insert a hard page break paragraph."""
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_break(
            # WD_BREAK.PAGE = 0  — use the integer to avoid import clutter
            __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
        )
    def export_layout(
        self, blocks: list[LayoutBlock], output_path: str
    ) -> str:
        """
        Write a DOCX from a list of LayoutBlocks.

        This is the primary export path from Milestone 8 onwards.
        Uses heading styles, correct spacing, bold, italic, and
        alignment from the detected layout.

        Parameters
        ----------
        blocks      : From LayoutDetector.detect()
        output_path : Full path for the output .docx file

        Returns
        -------
        str : Resolved output path
        """
        if not blocks:
            raise ValueError("No layout blocks to export.")

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        _set_document_rtl(doc)

        current_page = None

        for block in blocks:
            # Insert page break when page changes
            if current_page is not None and block.page_number != current_page:
                self._add_page_break(doc)
            current_page = block.page_number

            # Skip page numbers — they're noise in a reflowed DOCX
            if block.is_page_number:
                continue
            elif block.is_image:
                self._write_image_block(doc, block)
            elif block.is_table:
                self._write_table_block(doc, block)
            elif block.is_heading:
                self._write_layout_heading(doc, block)
            else:
                self._write_layout_paragraph(doc, block)
                
        try:
            doc.save(str(out))
            logger.info("DOCX (layout) saved: %s", out)
        except Exception as exc:
            raise IOError(f"Could not save DOCX: {exc}") from exc

        return str(out)

    def _write_layout_heading(self, doc: Document, block: LayoutBlock) -> None:
        """
        Write a heading block, reproducing any background color or border
        detected from the source PDF's vector graphics.
        """
        level = max(1, min(block.heading_level or 2, 3))
        para = doc.add_paragraph()
        _set_paragraph_rtl(para)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(4)
        # Apply detected background color (the actual box color from PDF)
        if block.background_color:
            self._apply_paragraph_shading(para, block.background_color)
        # Apply detected border
        if block.border_color and block.border_width > 0:
            self._apply_paragraph_border(para, block.border_color, block.border_width)
        run = para.add_run(block.text)
        run.font.name = self.font_name
        run.font.bold = True
        run.font.size = Pt(
            block.font_size if block.font_size
            else {1: 18, 2: 15, 3: 13}.get(level, 13)
        )
        # Use detected text color, or auto-pick based on background luminance
        if block.text_color:
            run.font.color.rgb = RGBColor(*block.text_color)
        elif block.background_color:
            run.font.color.rgb = self._contrast_text_color(block.background_color)
        _set_run_rtl(run)
    def _write_layout_paragraph(self, doc: Document, block: LayoutBlock) -> None:
        """
        Write a body paragraph, reproducing background and border from PDF.
        """
        para = doc.add_paragraph()
        _set_paragraph_rtl(para)

        if block.space_before > 0:
            para.paragraph_format.space_before = Pt(min(block.space_before, 24))

        if not block.is_rtl and block.alignment == Alignment.LEFT:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif not block.is_rtl and block.alignment == Alignment.CENTER:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply detected background and border
        if block.background_color:
            self._apply_paragraph_shading(para, block.background_color)
        if block.border_color and block.border_width > 0:
            self._apply_paragraph_border(para, block.border_color, block.border_width)

        run = para.add_run(block.text)
        run.font.name   = self.font_name
        run.font.size   = Pt(block.font_size) if block.font_size else self.font_size
        run.font.bold   = block.is_bold
        run.font.italic = block.is_italic

        if block.text_color:
            run.font.color.rgb = RGBColor(*block.text_color)
        elif block.background_color:
            run.font.color.rgb = self._contrast_text_color(block.background_color)

        _set_run_rtl(run)

    # ------------------------------------------------------------------
    # OOXML helpers for visual decoration
    # ------------------------------------------------------------------

    def _rgb_to_hex(self, rgb: tuple[int, int, int]) -> str:
        """Convert (r, g, b) 0-255 tuple to OOXML hex string e.g. 'FF0000'."""
        return "{:02X}{:02X}{:02X}".format(*rgb)

    def _apply_paragraph_shading(self, paragraph, rgb: tuple[int, int, int]) -> None:
        """
        Apply a background fill color to a paragraph using OOXML <w:shd>.
        This is the only way to set paragraph background in python-docx.
        """
        pPr = paragraph._p.get_or_add_pPr()
        # Remove existing shading
        for existing in pPr.findall(qn("w:shd")):
            pPr.remove(existing)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  self._rgb_to_hex(rgb))
        pPr.append(shd)

    def _apply_paragraph_border(
        self,
        paragraph,
        rgb:   tuple[int, int, int],
        width: float,
    ) -> None:
        """
        Apply borders to a paragraph using OOXML <w:pBdr>.
        Width is converted from PDF points to OOXML eighths-of-a-point.
        """
        pPr    = paragraph._p.get_or_add_pPr()
        pBdr   = OxmlElement("w:pBdr")
        hex_color = self._rgb_to_hex(rgb)
        # OOXML border size is in 1/8 pt; clamp between 2 and 18
        size = max(2, min(int(width * 8), 18))

        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    str(size))
            border.set(qn("w:space"), "4")
            border.set(qn("w:color"), hex_color)
            pBdr.append(border)

        pPr.append(pBdr)

    def _contrast_text_color(
        self, background_rgb: tuple[int, int, int]
    ) -> RGBColor:
        """
        Return white or black text depending on background luminance.
        Uses the WCAG relative luminance formula.
        """
        r, g, b = (c / 255.0 for c in background_rgb)
        luminance = 0.2126 * r + 0.7152 * g + 0.0722 * b
        return RGBColor(255, 255, 255) if luminance < 0.5 else RGBColor(0, 0, 0)

    def _write_image_block(self, doc: Document, block: LayoutBlock) -> None:
        """
        Embed an extracted image into the DOCX.

        Scales the image to fit within the page content width (6 inches)
        while preserving the original aspect ratio.
        """
        if not block.image_path or not Path(block.image_path).exists():
            logger.warning("Image file not found: %s — skipping.", block.image_path)
            return

        try:
            max_width_inches = 6.0

            if block.image_width and block.image_height and block.image_width > 0:
                # Convert points to inches (72 points per inch)
                width_inches  = min(block.image_width / 72, max_width_inches)
                aspect        = block.image_height / block.image_width
                height_inches = width_inches * aspect
                doc.add_picture(block.image_path, width=Inches(width_inches))
            else:
                doc.add_picture(block.image_path, width=Inches(max_width_inches))

            # Centre the image paragraph
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as exc:
            logger.warning("Could not embed image %s: %s", block.image_path, exc)
            # Add a text placeholder so the reader knows an image was here
            para = doc.add_paragraph()
            run  = para.add_run(f"[ Image ]")
            run.font.italic = True
            run.font.color.rgb = PLACEHOLDER_COLOR

    def _write_table_block(self, doc: Document, block: LayoutBlock) -> None:
        """
        Render a detected table into the DOCX.

        Creates a python-docx Table with one cell per detected cell.
        Applies light borders and RTL direction to every cell.
        """
        if not block.table_data:
            return

        rows    = len(block.table_data)
        cols    = max(len(row) for row in block.table_data)

        try:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"

            for r_idx, row_data in enumerate(block.table_data):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx >= cols:
                        break
                    cell      = table.cell(r_idx, c_idx)
                    cell.text = cell_text.strip()

                    # Apply RTL to cell paragraphs
                    for para in cell.paragraphs:
                        _set_paragraph_rtl(para)
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in para.runs:
                            run.font.name = self.font_name
                            run.font.size = Pt(11)
                            _set_run_rtl(run)

        except Exception as exc:
            logger.warning("Could not write table: %s", exc)
            # Fallback: write cells as tab-separated text
            para = doc.add_paragraph()
            for row in block.table_data:
                para.add_run(" | ".join(cell for cell in row) + "\n")