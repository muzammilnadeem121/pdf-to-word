"""
Tests for LayoutDetector and LayoutBlock models.

No real PDF required for most tests — we use synthetic pdfplumber-style
word dicts and ExtractionResult fixtures.

Run with:
    python -m pytest tests/test_layout.py -v
"""

import pytest
from pathlib import Path

from models.layout_block import LayoutBlock, BlockType, Alignment
from layout.detector import (
    LayoutDetector,
    _detect_alignment,
    _is_rtl_text,
)
from extractor.extractor import ExtractionResult, PageResult
from models.page_classification import PageClassification, PageType


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def make_classification(page_number=1, page_type=PageType.DIGITAL) -> PageClassification:
    return PageClassification(
        page_number     = page_number,
        page_type       = page_type,
        confidence      = 0.9,
        char_count      = 50,
        image_coverage  = 0.0,
        urdu_char_ratio = 0.8,
        reason          = "test",
    )


def make_page_result(
    page_number = 1,
    page_type   = PageType.DIGITAL,
    raw_text    = "یہ ایک آزمائشی متن ہے جو کافی لمبا ہے",
) -> PageResult:
    return PageResult(
        page_number    = page_number,
        classification = make_classification(page_number, page_type),
        raw_text       = raw_text,
        char_count     = len(raw_text.replace(" ", "")),
    )


def make_extraction(pages: list[PageResult]) -> ExtractionResult:
    return ExtractionResult(
        file_path     = "dummy.pdf",
        total_pages   = len(pages),
        digital_pages = sum(1 for p in pages if p.is_digital),
        scanned_pages = sum(1 for p in pages if p.is_scanned),
        mixed_pages   = sum(1 for p in pages if p.is_mixed),
        pages         = pages,
    )


# ---------------------------------------------------------------------------
# LayoutBlock model
# ---------------------------------------------------------------------------

class TestLayoutBlock:

    def test_is_heading_true(self):
        block = LayoutBlock(
            text="عنوان", block_type=BlockType.HEADING,
            page_number=1, heading_level=1,
        )
        assert block.is_heading is True
        assert block.is_body is False

    def test_is_body_true(self):
        block = LayoutBlock(
            text="متن", block_type=BlockType.BODY, page_number=1
        )
        assert block.is_body is True
        assert block.is_heading is False

    def test_is_page_number_true(self):
        block = LayoutBlock(
            text="1", block_type=BlockType.PAGE_NUM, page_number=1
        )
        assert block.is_page_number is True

    def test_default_alignment_is_right(self):
        block = LayoutBlock(text="متن", block_type=BlockType.BODY, page_number=1)
        assert block.alignment == Alignment.RIGHT

    def test_default_is_rtl(self):
        block = LayoutBlock(text="متن", block_type=BlockType.BODY, page_number=1)
        assert block.is_rtl is True


# ---------------------------------------------------------------------------
# Alignment detection
# ---------------------------------------------------------------------------

class TestAlignmentDetection:

    def test_right_aligned(self):
        # Block flush to right margin (x1 close to page width)
        result = _detect_alignment(x0=300, x1=500, page_width=520)
        assert result == Alignment.RIGHT

    def test_centred(self):
        # Block centred on a 600-wide page (centre=300, block centre=260–340)
        result = _detect_alignment(x0=200, x1=400, page_width=600)
        assert result == Alignment.CENTER

    def test_left_aligned(self):
        result = _detect_alignment(x0=10, x1=200, page_width=600)
        assert result == Alignment.LEFT

    def test_zero_page_width_returns_right(self):
        result = _detect_alignment(x0=0, x1=100, page_width=0)
        assert result == Alignment.RIGHT


# ---------------------------------------------------------------------------
# RTL text detection
# ---------------------------------------------------------------------------

class TestRTLDetection:

    def test_urdu_text_is_rtl(self):
        assert _is_rtl_text("نماز پڑھنا") is True

    def test_english_text_not_rtl(self):
        assert _is_rtl_text("Hello world") is False

    def test_mixed_mostly_urdu_is_rtl(self):
        assert _is_rtl_text("نماز abc") is True

    def test_empty_string_not_rtl(self):
        assert _is_rtl_text("") is False


# ---------------------------------------------------------------------------
# Classification logic
# ---------------------------------------------------------------------------

class TestClassification:

    def setup_method(self):
        self.detector = LayoutDetector()

    def test_large_font_classified_as_heading(self):
        block_type, level = self.detector._classify_line(
            text="عنوان", font_size=24.0, body_size=12.0,
            is_bold=False, top=100, page_height=800,
        )
        assert block_type == BlockType.HEADING
        assert level is not None

    def test_body_size_classified_as_body(self):
        block_type, level = self.detector._classify_line(
            text="یہ ایک لمبا جملہ ہے جو عام متن کی طرح ہے",
            font_size=12.0, body_size=12.0,
            is_bold=False, top=400, page_height=800,
        )
        assert block_type == BlockType.BODY
        assert level is None

    def test_bold_short_line_is_heading(self):
        block_type, level = self.detector._classify_line(
            text="عنوان فہرست", font_size=12.0, body_size=12.0,
            is_bold=True, top=400, page_height=800,
        )
        assert block_type == BlockType.HEADING

    def test_numeric_in_margin_is_page_number(self):
        block_type, level = self.detector._classify_line(
            text="١٢", font_size=10.0, body_size=12.0,
            is_bold=False, top=5, page_height=800,  # near top margin
        )
        assert block_type == BlockType.PAGE_NUM

    def test_heading_level_1_very_large(self):
        level = self.detector._heading_level(font_size=28.0, body_size=12.0)
        assert level == 1

    def test_heading_level_2_moderately_large(self):
        level = self.detector._heading_level(font_size=20.0, body_size=12.0)
        assert level == 2

    def test_heading_level_3_slightly_large(self):
        level = self.detector._heading_level(font_size=15.0, body_size=12.0)
        assert level == 3


# ---------------------------------------------------------------------------
# Line grouping
# ---------------------------------------------------------------------------

class TestLineGrouping:

    def setup_method(self):
        self.detector = LayoutDetector()

    def _word(self, text, top, x0, x1) -> dict:
        return {"text": text, "top": top, "x0": x0, "x1": x1,
                "size": 12.0, "fontname": "Arial"}

    def test_same_top_grouped_into_one_line(self):
        words = [
            self._word("نماز", top=100, x0=300, x1=350),
            self._word("پڑھنا", top=100, x0=200, x1=290),
        ]
        lines = self.detector._group_words_into_lines(words)
        assert len(lines) == 1
        assert len(lines[0]) == 2

    def test_different_tops_form_separate_lines(self):
        words = [
            self._word("پہلی", top=100, x0=300, x1=350),
            self._word("دوسری", top=130, x0=300, x1=360),
        ]
        lines = self.detector._group_words_into_lines(words, tolerance=3.0)
        assert len(lines) == 2

    def test_within_tolerance_grouped_together(self):
        words = [
            self._word("الف", top=100, x0=300, x1=330),
            self._word("ب",   top=102, x0=200, x1=220),  # 2pt diff, within 3pt
        ]
        lines = self.detector._group_words_into_lines(words, tolerance=3.0)
        assert len(lines) == 1

    def test_empty_words_returns_empty(self):
        assert self.detector._group_words_into_lines([]) == []


# ---------------------------------------------------------------------------
# Fallback plain text
# ---------------------------------------------------------------------------

class TestFallback:

    def test_fallback_produces_body_blocks(self):
        page   = make_page_result(raw_text="نماز پڑھنا\nروزہ رکھنا")
        result = make_extraction([page])
        detector = LayoutDetector()
        blocks = detector._fallback_plain_text(result)
        assert len(blocks) == 2
        assert all(b.block_type == BlockType.BODY for b in blocks)

    def test_fallback_empty_text_returns_empty(self):
        page   = make_page_result(raw_text="")
        result = make_extraction([page])
        detector = LayoutDetector()
        blocks = detector._fallback_plain_text(result)
        assert blocks == []

    def test_fallback_preserves_page_numbers(self):
        page   = make_page_result(page_number=3, raw_text="متن")
        result = make_extraction([page])
        blocks = LayoutDetector()._fallback_plain_text(result)
        assert blocks[0].page_number == 3


# ---------------------------------------------------------------------------
# WordExporter.export_layout
# ---------------------------------------------------------------------------

class TestExportLayout:

    def _make_blocks(self) -> list[LayoutBlock]:
        return [
            LayoutBlock(
                text="عنوان", block_type=BlockType.HEADING,
                page_number=1, heading_level=1, is_rtl=True,
            ),
            LayoutBlock(
                text="یہ ایک لمبا جملہ ہے جو عام متن کی طرح ہے",
                block_type=BlockType.BODY,
                page_number=1, is_rtl=True,
            ),
            LayoutBlock(
                text="1", block_type=BlockType.PAGE_NUM, page_number=1,
            ),
            LayoutBlock(
                text="دوسرا صفحہ", block_type=BlockType.BODY,
                page_number=2, is_rtl=True,
            ),
        ]

    def test_creates_docx_file(self, tmp_path):
        from exporter.word_export import WordExporter
        blocks = self._make_blocks()
        out    = tmp_path / "layout.docx"
        WordExporter().export_layout(blocks, str(out))
        assert out.exists()

    def test_raises_on_empty_blocks(self, tmp_path):
        from exporter.word_export import WordExporter
        with pytest.raises(ValueError):
            WordExporter().export_layout([], str(tmp_path / "empty.docx"))

    def test_page_numbers_skipped(self, tmp_path):
        from exporter.word_export import WordExporter
        from docx import Document
        blocks = self._make_blocks()
        out    = tmp_path / "no_pagenum.docx"
        WordExporter().export_layout(blocks, str(out))
        doc       = Document(str(out))
        full_text = " ".join(p.text for p in doc.paragraphs)
        # "1" (the page number block) should not appear as a standalone paragraph
        assert full_text.count("1") == 0

    def test_bold_block_written(self, tmp_path):
        from exporter.word_export import WordExporter
        from docx import Document
        blocks = [
            LayoutBlock(
                text="굵은 متن", block_type=BlockType.BODY,
                page_number=1, is_bold=True, is_rtl=True,
            )
        ]
        out = tmp_path / "bold.docx"
        WordExporter().export_layout(blocks, str(out))
        doc  = Document(str(out))
        runs = [r for p in doc.paragraphs for r in p.runs]
        assert any(r.bold for r in runs)